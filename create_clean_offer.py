from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/home/ubuntu/boostnow/oferta_self_storage.docx"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, (255,255,255))
        shade(t.rows[0].cells[i], '1F2937')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
styles['Normal'].paragraph_format.space_after = Pt(5)
for name in ['Title', 'Heading 1', 'Heading 2']:
    styles[name].font.name = 'Arial'
    styles[name].font.color.rgb = RGBColor(31,41,55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OFERTA USŁUG')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(31,41,55)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('marketing cyfrowy dla firm self-storage')
r.font.name = 'Arial'
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(75,85,99)
p.paragraph_format.space_after = Pt(15)

p = doc.add_paragraph()
p.add_run('Zakres dokumentu: ').bold = True
p.add_run('obsługa komunikacji firmy, profilu Google, treści, krótkich materiałów wideo, kampanii reklamowych i raportowania. Ceny podano w PLN netto.')

heading(doc, '1. Pakiety miesięczne', 1)
table(doc, ['Pakiet', 'Cena netto / mies.', 'Zakres podstawowy'], [
    ['Podstawowy', '1 490 zł', 'Profil firmy w Google, 8 postów, 2 rolki, 4 relacje, raport miesięczny'],
    ['Rozszerzony', '2 490 zł', 'Google, Instagram i Facebook, 12 postów, 4 rolki, 8 relacji, Meta Ads, pomiar leadów'],
    ['Pełna obsługa', 'od 3 990 zł', '18 postów, 6 rolek, 12 relacji, kampanie reklamowe, SEO firmy, raport zbiorczy'],
], [1.35,1.15,4.95])

heading(doc, '2. Zakres prac', 1)
table(doc, ['Element', 'Pakiet podstawowy', 'Pakiet rozszerzony', 'Pakiet pełny'], [
    ['Profil firmy w Google', 'optymalizacja i publikacje', 'optymalizacja, publikacje i pomiar', 'obsługa profilu i standard komunikacji'],
    ['Social media', '1 kanał', 'Instagram + Facebook', 'Instagram + Facebook'],
    ['Posty', '8 / mies.', '12 / mies.', '18 / mies. łącznie'],
    ['Rolki', '2 / mies.', '4 / mies.', '6 / mies.'],
    ['Relacje', '4 / mies.', '8 / mies.', '12 / mies.'],
    ['SEO / strony', '1 temat / mies.', '2 tematy / mies. lub landing page co 2 mies.', '1 optymalizacja lub landing page / mies.'],
    ['Reklamy', 'brak', '1–2 kampanie Meta', 'kampanie Meta i Google według ustaleń'],
    ['Moderacja', 'do 1 godz. / mies.', 'do 2 godz. / mies.', 'ustalana w zakresie projektu'],
    ['Raportowanie', 'raport miesięczny', 'raport leadów i kosztu leada', 'dashboard i porównanie wyników'],
], [1.55,1.65,2.1,2.15])

heading(doc, '3. Usługi jednorazowe i dodatki', 1)
table(doc, ['Usługa', 'Cena netto', 'Zakres'], [
    ['Onboarding firmy', '1 490 zł', 'audyt Google, social media, strony, konkurencji i konfiguracja pomiaru'],
    ['Start komunikacji firmy', '3 900 zł', 'przygotowanie profili, landing page, kampania startowa i treści na 30 dni'],
    ['Pakiet 8 rolek', '2 400 zł', 'montaż z materiałów dostarczonych przez klienta, napisy i okładki'],
    ['Landing page firmy', 'od 1 900 zł', 'strona oferty firmy z CTA, formularzem i analityką'],
    ['Audyt Google Business Profile', '790 zł', 'kategorie, zdjęcia, opinie, profil firmy i rekomendacje'],
    ['Prowadzenie Google Ads', 'od 900 zł / mies.', 'obsługa kampanii; budżet reklamowy rozliczany osobno'],
    ['Dodatkowe 4 posty', '500 zł / mies.', 'przygotowanie, montaż i publikacja dodatkowych materiałów'],
    ['Dodatkowy kanał social', '350 zł / mies.', 'adaptacja i publikacja istniejącego materiału'],
    ['Dodatkowe 4 rolki', '700 zł / mies.', 'montaż z materiałów dostarczonych przez klienta'],
], [2.0,1.35,4.1])

heading(doc, '4. Warunki współpracy', 1)
for text in [
    'Ceny są cenami netto i nie obejmują budżetu reklamowego ani kosztów narzędzi zewnętrznych. Zdjęcia i nagrania dostarcza klient.',
    'Minimalny okres współpracy wynosi 3 miesiące. Pierwszy miesiąc obejmuje konfigurację, przygotowanie materiałów i zebranie danych.',
    'Onboarding jest rozliczany jednorazowo przed rozpoczęciem obsługi miesięcznej.',
    'Budżet reklamowy pozostaje własnością klienta i jest opłacany bezpośrednio na platformie reklamowej.',
    'Klient dostarcza zdjęcia, nagrania, informacje, dostępy i akceptacje. Usługa obejmuje selekcję, montaż, przygotowanie publikacji i publikację materiałów.',
    'Oferta nie obejmuje gwarancji określonej liczby leadów, rezerwacji ani pozycji w wynikach wyszukiwania.',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

heading(doc, '5. Zakres rozszerzony', 1)
doc.add_paragraph('Zakres może zostać rozszerzony o dodatkowe kanały, większą liczbę publikacji, większą liczbę kampanii lub dodatkowe formaty materiałów. Wycena zależy od liczby materiałów i częstotliwości publikacji.')

table(doc, ['Przykładowy zakres', 'Cena netto / mies.'], [
    ['Dodatkowy kanał komunikacji', '350 zł'],
    ['Dodatkowe 4 rolki', '700 zł'],
    ['Indywidualny zakres publikacji i kampanii', 'wycena indywidualna'],
], [5.2,2.2])

heading(doc, '6. Dane do rozpoczęcia współpracy', 1)
doc.add_paragraph('Do rozpoczęcia współpracy potrzebne są: nazwa i opis firmy, oferta usług, grupa klientów, dostępne zdjęcia i nagrania dostarczone przez klienta, dostępy do profilu Google oraz kanałów social media, a także osoba odpowiedzialna za przekazywanie materiałów, akceptacje i obsługę zapytań.')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ważność oferty: 14 dni od daty przekazania. Wszystkie kwoty netto.')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(107,114,128)

doc.save(OUT)
print(OUT)
