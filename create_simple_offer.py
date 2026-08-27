from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = '/home/ubuntu/boostnow/oferta_self_storage.docx'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_text(cell, text, bold=False, white=False, size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def make_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_text(t.rows[0].cells[i], header, bold=True, white=True, size=8)
        shade(t.rows[0].cells[i], '263238')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_text(cells[i], value, size=8)
    for row in t.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    return t

def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(38, 50, 56)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.45)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.55)
sec.right_margin = Inches(0.55)

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(8.5)
doc.styles['Normal'].paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run('OFERTA OBSŁUGI MEDIÓW SPOŁECZNOŚCIOWYCH')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(38, 50, 56)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('dla firm self-storage')
r.font.name = 'Arial'
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(84, 110, 122)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('Jak działa współpraca: ').bold = True
p.add_run('klient przekazuje zdjęcia, nagrania i informacje. W ramach pakietu przygotowujemy treści, montujemy rolki, dodajemy napisy, publikujemy materiały i przekazujemy krótkie podsumowanie działań.')

section_title(doc, 'PAKIETY MIESIĘCZNE')
make_table(doc, ['Pakiet', 'Cena netto', 'W pakiecie'], [
    ['Podstawowy', '1 490 zł / mies.', '8 postów • 2 rolki • 4 relacje • 1 kanał social • publikacja'],
    ['Rozszerzony', '2 490 zł / mies.', '12 postów • 4 rolki • 8 relacji • Instagram + Facebook • profil Google • raport'],
    ['Pełny', '3 490 zł / mies.', '16 postów • 6 rolek • 12 relacji • Instagram + Facebook • profil Google • obsługa kampanii Meta • raport'],
], [1.25,1.45,5.0])

section_title(doc, 'CO OZNACZAJĄ ELEMENTY PAKIETU?')
make_table(doc, ['Element', 'Zakres'], [
    ['Post', 'Przygotowanie tekstu, grafiki lub obróbka dostarczonego materiału oraz publikacja.'],
    ['Rolka', 'Montaż dostarczonego nagrania, napisy, muzyka z biblioteki i przygotowanie do publikacji.'],
    ['Relacja', 'Krótki materiał do Stories; zdjęcie, grafika lub fragment nagrania dostarczony przez klienta.'],
    ['Raport', 'Liczba publikacji, podstawowe statystyki kanałów i najważniejsze obserwacje.'],
], [1.0,6.7])

section_title(doc, 'DODATKOWO, JEŚLI POTRZEBNE')
make_table(doc, ['Usługa', 'Cena netto'], [
    ['Dodatkowe 4 rolki z materiałów klienta', '700 zł / mies.'],
    ['Dodatkowe 4 posty', '500 zł / mies.'],
    ['Dodatkowy kanał social', '350 zł / mies.'],
    ['Prowadzenie Google Ads', 'od 900 zł / mies.'],
], [6.1,1.6])

section_title(doc, 'WARUNKI')
for text in [
    'Ceny są netto. Budżet reklamowy jest płatny osobno bezpośrednio do platformy reklamowej.',
    'Klient dostarcza zdjęcia, nagrania, informacje, dostępy i akceptacje. Usługa nie obejmuje nagrywania materiałów.',
    'Minimalny okres współpracy: 3 miesiące. Pierwszy miesiąc obejmuje konfigurację i przygotowanie procesu publikacji.',
    'Termin publikacji zależy od terminowego przekazania materiałów i akceptacji po stronie klienta.',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.18)
    p.add_run(text)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(5)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ważność oferty: 14 dni • Wszystkie kwoty netto')
r.italic = True
r.font.name = 'Arial'
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(84, 110, 122)

doc.save(OUT)
print(OUT)
